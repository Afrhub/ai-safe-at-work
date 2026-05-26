#!/usr/bin/env python3
"""
Export all course content to a single Word document.

Reads every HTML page from the project root + templates/, strips
chrome (topbar, footer, JSON-LD, style blocks), and renders structured
content (h1-h4, paragraphs, lists, tables, callouts, code) into one
clean .docx for PO review.

Output: ../AI-Safe-At-Work-content.docx
"""
import re
import sys
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT.parent / "AI-Safe-At-Work-content.docx"

# Ordered list of pages — controls reading order in the document.
PAGES = [
    ("Cover · Landing page", "index.html"),
    ("Course Overview · all modules listed", "course.html"),

    ("Module 1 — Why this course exists", "module-1.html"),
    ("Module 2 — What AI tools do with what you type", "module-2.html"),
    ("Module 3 — The never-paste list", "module-3.html"),
    ("Module 4 — Picking the right tool for the job", "module-4.html"),
    ("Module 5 — Verifying what the AI tells you", "module-5.html"),
    ("Module 6 — AI-powered scams aimed at you", "module-6.html"),
    ("Module 7 — Bias, fairness, and not embarrassing the business", "module-7.html"),
    ("Module 8 — Copyright, IP, and other people's content", "module-8.html"),
    ("Module 9 — Logging, accountability, and 'who used what'", "module-9.html"),
    ("Module 10 — When something goes wrong", "module-10.html"),
    ("Module 11 — The 60-second pre-submit checklist", "module-11.html"),
    ("Module 12 — The standards behind this course", "module-12.html"),

    ("Role track — Manager / Leadership", "module-manager.html"),
    ("Role track — DPO / Compliance", "module-dpo.html"),

    ("Standards Map (clause-to-module)", "standards-map.html"),
    ("Citations bibliography", "citations.html"),
    ("Changelog", "changelog.html"),

    ("Template — Acceptable Use Policy", "templates/aup-template.html"),
    ("Template — AI vendor diligence questionnaire", "templates/vendor-questionnaire.html"),

    ("MSP partner programme (draft)", "msp.html"),
    ("Pricing (draft)", "pricing.html"),
    ("Security & AI-crawler posture", "security.html"),

    ("Accessibility statement", "accessibility.html"),
    ("Complaints procedure", "complaints.html"),
    ("Privacy notice", "privacy.html"),
    ("Terms of use", "terms.html"),
]

ACCENT = RGBColor(0xB6, 0x77, 0x00)         # amber for headings
MUTED = RGBColor(0x55, 0x55, 0x55)          # grey for callout-icon column
TEXT = RGBColor(0x18, 0x18, 0x18)           # near-black body
TABLE_HEAD = RGBColor(0xB6, 0x77, 0x00)

def add_horizontal_line(paragraph):
    """Add a thin grey horizontal line beneath a paragraph (used for h2 separator)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B67700")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def text_of(node):
    """Return the visible text from a tag, normalising whitespace."""
    if node is None:
        return ""
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True)).strip()

def add_styled_run(paragraph, text, *, bold=False, italic=False, color=None, size=None, mono=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if mono:
        run.font.name = "Consolas"
    return run

def render_inline(paragraph, node):
    """Walk inline children of a tag and add styled runs."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                add_styled_run(paragraph, re.sub(r"\s+", " ", text))
            elif text:
                add_styled_run(paragraph, " ")
        elif isinstance(child, Tag):
            tname = child.name.lower()
            if tname in ("strong", "b"):
                add_styled_run(paragraph, text_of(child), bold=True)
            elif tname in ("em", "i"):
                add_styled_run(paragraph, text_of(child), italic=True, color=ACCENT)
            elif tname == "code":
                add_styled_run(paragraph, text_of(child), mono=True, color=ACCENT, size=10)
            elif tname == "a":
                href = child.get("href", "")
                txt = text_of(child)
                if href and href.startswith("http"):
                    add_styled_run(paragraph, f"{txt} ({href})", color=ACCENT)
                else:
                    add_styled_run(paragraph, txt, color=ACCENT)
            elif tname == "br":
                add_styled_run(paragraph, "\n")
            elif tname == "span":
                # Recurse into spans (callouts use them)
                render_inline(paragraph, child)
            else:
                # Fallback — plain text
                add_styled_run(paragraph, text_of(child))

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    if level == 1:
        p_format.space_before = Pt(28)
        p_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        add_horizontal_line(p)
    elif level == 2:
        p_format.space_before = Pt(18)
        p_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = TEXT
    elif level == 3:
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = TEXT
    else:  # h4+
        p_format.space_before = Pt(8)
        p_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEXT
    return p

def add_paragraph(doc, node):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    render_inline(p, node)
    for run in p.runs:
        if run.font.size is None:
            run.font.size = Pt(11)
        if run.font.color.rgb is None:
            run.font.color.rgb = TEXT
    return p

def add_list(doc, node, *, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.add_run("• " if not ordered else "1. ")
        p.paragraph_format.space_after = Pt(3)
        render_inline(p, li)
        # nested lists
        for sub in li.find_all(["ul", "ol"], recursive=False):
            sub_ordered = sub.name == "ol"
            for sub_li in sub.find_all("li", recursive=False):
                try:
                    sp = doc.add_paragraph(style="List Bullet 2")
                except KeyError:
                    sp = doc.add_paragraph()
                    sp.paragraph_format.left_indent = Cm(1.6)
                    sp.add_run("◦ " if not sub_ordered else "  • ")
                sp.paragraph_format.space_after = Pt(2)
                render_inline(sp, sub_li)
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(11)
            if run.font.color.rgb is None:
                run.font.color.rgb = TEXT

def add_table_from_html(doc, table_node):
    rows = table_node.find_all("tr")
    if not rows:
        return
    # Determine column count from longest row
    col_count = max(len(r.find_all(["th", "td"])) for r in rows)
    docx_table = doc.add_table(rows=len(rows), cols=col_count)
    docx_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        docx_table.style = "Light Grid Accent 1"
    except KeyError:
        docx_table.style = None
    for r_i, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        is_header = all(c.name == "th" for c in cells) if cells else False
        for c_i in range(col_count):
            cell = docx_table.cell(r_i, c_i)
            cell.text = ""
            if c_i < len(cells):
                src = cells[c_i]
                p = cell.paragraphs[0]
                render_inline(p, src)
                for run in p.runs:
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = TABLE_HEAD
                    if run.font.size is None:
                        run.font.size = Pt(10)
            if is_header:
                cell_shade(cell, "FFF5DD")
    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_callout(doc, callout):
    """Render a .callout div as an indented styled paragraph."""
    title_node = callout.find(class_="callout-title")
    title = text_of(title_node) if title_node else ""
    classes = callout.get("class", [])
    kind = next((c for c in classes if c in ("warn", "do", "dont", "note", "standard")), "note")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    label_map = {"warn": "WARNING", "do": "DO", "dont": "DON'T", "standard": "STANDARDS", "note": "NOTE"}
    add_styled_run(p, label_map.get(kind, "NOTE") + " · ", bold=True, color=ACCENT, size=9)
    if title:
        add_styled_run(p, title, bold=True, size=11, color=TEXT)

    # Body content — find inner div, render its children
    inner = callout.find("div")
    if inner is not None:
        # remove callout-title from inner to avoid duplicate
        for ct in inner.find_all(class_="callout-title"):
            ct.decompose()
        for child in inner.children:
            if isinstance(child, Tag):
                tname = child.name.lower()
                if tname == "p":
                    add_paragraph(doc, child)
                elif tname == "ul":
                    add_list(doc, child, ordered=False)
                elif tname == "ol":
                    add_list(doc, child, ordered=True)

def process_node(doc, node):
    """Walk one top-level node from <main> and render."""
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            p = doc.add_paragraph()
            add_styled_run(p, text, size=11, color=TEXT)
        return
    if not isinstance(node, Tag):
        return
    classes = node.get("class", [])
    tname = node.name.lower()

    # Skip noise
    if "no-print" in classes:
        return
    if tname in ("script", "style", "link", "meta", "header", "footer", "nav"):
        return

    if tname in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tname[1])
        add_heading(doc, text_of(node), level)
    elif tname == "p":
        # Skip class 'eyebrow' and 'muted' subtle text in cleaner export — keep eyebrow as small caps
        if "eyebrow" in classes:
            p = doc.add_paragraph()
            add_styled_run(p, text_of(node).upper(), bold=True, color=ACCENT, size=9)
            p.paragraph_format.space_after = Pt(2)
        elif "lede" in classes:
            p = doc.add_paragraph()
            add_styled_run(p, text_of(node), size=12, color=TEXT)
            p.paragraph_format.space_after = Pt(8)
        else:
            add_paragraph(doc, node)
    elif tname == "ul":
        add_list(doc, node, ordered=False)
    elif tname == "ol":
        add_list(doc, node, ordered=True)
    elif tname == "hr":
        # Section divider — add small blank space
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)
        add_horizontal_line(sp)
    elif tname == "table":
        add_table_from_html(doc, node)
    elif "callout" in classes:
        add_callout(doc, node)
    elif "objectives" in classes:
        # Render objectives block as a styled callout
        title_node = node.find(class_="objectives-title")
        if title_node:
            p = doc.add_paragraph()
            add_styled_run(p, text_of(title_node).upper(), bold=True, color=ACCENT, size=10)
            p.paragraph_format.space_after = Pt(4)
        for ul in node.find_all("ul", recursive=False):
            add_list(doc, ul, ordered=False)
    elif "compare" in classes:
        # Compare grid → 2-col table
        cells = node.find_all(class_="compare-cell")
        if cells:
            tbl = doc.add_table(rows=1, cols=len(cells))
            try:
                tbl.style = "Light Grid Accent 1"
            except KeyError:
                pass
            for i, c in enumerate(cells):
                cell = tbl.cell(0, i)
                cell.text = ""
                for ch in c.children:
                    if isinstance(ch, Tag):
                        if ch.name == "h4":
                            p = cell.add_paragraph()
                            add_styled_run(p, text_of(ch), bold=True, size=11, color=TABLE_HEAD)
                        elif ch.name == "p":
                            p = cell.add_paragraph()
                            render_inline(p, ch)
                            for r in p.runs:
                                if r.font.size is None:
                                    r.font.size = Pt(10)
            doc.add_paragraph()
    elif "checklist" in classes:
        # Render check items as bullet list with bold strong
        for item in node.find_all(class_="check-item"):
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("• ")
            strong = item.find("strong")
            span = item.find("span")
            if strong:
                add_styled_run(p, text_of(strong) + " — ", bold=True, color=TEXT)
            if span:
                add_styled_run(p, text_of(span))
    elif "module-grid" in classes:
        # Course overview cards → numbered list with title + description
        for card in node.find_all(class_="module-card", recursive=False):
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            num = card.find(class_="module-num")
            h3 = card.find("h3")
            pp = card.find("p")
            label = ""
            if num: label += text_of(num) + " · "
            if h3:  label += text_of(h3)
            add_styled_run(p, label, bold=True, color=TEXT)
            if pp:
                ip = doc.add_paragraph()
                ip.paragraph_format.left_indent = Cm(0.8)
                add_styled_run(ip, text_of(pp), size=10, color=MUTED)
    elif "tier-grid" in classes:
        # Pricing tier cards → table
        cards = node.find_all(class_="tier", recursive=False)
        if cards:
            tbl = doc.add_table(rows=4, cols=len(cards))
            try:
                tbl.style = "Light Grid Accent 1"
            except KeyError:
                pass
            for i, card in enumerate(cards):
                tag = card.find(class_="tier-tag")
                h3 = card.find("h3")
                price = card.find(class_="price")
                ul = card.find("ul")
                tbl.cell(0, i).text = text_of(tag) if tag else ""
                tbl.cell(1, i).text = text_of(h3) if h3 else ""
                tbl.cell(2, i).text = text_of(price) if price else ""
                feats_cell = tbl.cell(3, i)
                feats_cell.text = ""
                if ul:
                    for li in ul.find_all("li", recursive=False):
                        fp = feats_cell.add_paragraph()
                        add_styled_run(fp, "• " + text_of(li), size=9)
            doc.add_paragraph()
    elif tname in ("div", "section", "article", "main", "aside"):
        # Recurse into containers
        for child in node.children:
            process_node(doc, child)
    elif tname == "details":
        # Render details block
        for child in node.children:
            process_node(doc, child)
    # Anything else — ignore (script, link, etc.)

def render_page(doc, title, html_path):
    """Render one HTML page as a Word section."""
    html = html_path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "lxml")

    # Page break before each section except first
    if doc.element.body.findall(qn("w:p")):
        doc.add_page_break()

    # Section title
    add_heading(doc, title, level=1)
    rel = str(html_path.relative_to(ROOT)).replace("\\", "/")
    src_p = doc.add_paragraph()
    add_styled_run(src_p, f"Source · {rel}", italic=True, color=MUTED, size=9)
    src_p.paragraph_format.space_after = Pt(12)

    # Main body — only inside <main>
    main = soup.find("main")
    if main is None:
        main = soup.body or soup
    for child in main.children:
        process_node(doc, child)

def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run("AI Safe@Work")
    run.font.size = Pt(36); run.font.bold = True; run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI Governance Awareness + Enablement Platform")
    run.font.size = Pt(16); run.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("Full content export · for PO review")
    run.font.size = Pt(11); run.font.italic = True; run.font.color.rgb = MUTED

    meta = [
        ("Version", "v2026.05"),
        ("Export date", "2026-05-19"),
        ("Live site", "https://aisafework.netlify.app/"),
        ("Repo", "https://github.com/Afrhub/ai-safe-at-work"),
        ("Licence (content)", "CC BY-SA 4.0"),
        ("Licence (code)", "MIT"),
        ("Standards", "EU AI Act · GDPR · ISO/IEC 42001 · ISO/IEC 27001 · NIST AI RMF · OWASP LLM Top 10"),
    ]
    for k, v in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_run(p, f"{k}: ", bold=True, color=ACCENT, size=10)
        add_styled_run(p, v, color=TEXT, size=10)

def add_toc(doc):
    doc.add_page_break()
    add_heading(doc, "Contents", level=1)
    for i, (title, _path) in enumerate(PAGES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        add_styled_run(p, f"{i:02d}.  ", bold=True, color=ACCENT, size=10)
        add_styled_run(p, title, size=10, color=TEXT)

def main():
    doc = Document()

    # Set default font + page margins
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_cover(doc)
    add_toc(doc)

    for title, rel_path in PAGES:
        path = ROOT / rel_path
        if not path.exists():
            print(f"  ! missing: {rel_path}", file=sys.stderr)
            continue
        print(f"  + {rel_path}")
        render_page(doc, title, path)

    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size:,} bytes")

if __name__ == "__main__":
    main()
